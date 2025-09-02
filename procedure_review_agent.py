"""
Quality Procedure Review Agent using LangGraph
Implements Evaluator-Optimizer pattern for ISO compliance review
"""

import os
from typing import TypedDict, List, Dict, Any
from langchain_openai import ChatOpenAI
from langchain_core.messages import HumanMessage, AIMessage
from langchain_core.prompts import ChatPromptTemplate
from langgraph.graph import StateGraph, END
from docx import Document
from docx.shared import Inches
from datetime import datetime
import json
from dotenv import load_dotenv

load_dotenv()

# State definition for the workflow
class ReviewState(TypedDict):
    input_procedure: str
    generated_review: str
    evaluation_result: Dict[str, Any]
    feedback: str
    final_output: str
    iteration_count: int
    max_iterations: int

class QualityReviewAgent:
    def __init__(self, openai_api_key: str):
        """Initialize the Quality Review Agent with OpenAI model"""
        self.llm = ChatOpenAI(
            api_key=openai_api_key,
            model="gpt-4o-mini",
            temperature=0.1
        )
        
        # ISO Standards Knowledge Base
        self.knowledge_base = {
            "ISO13485": {
                "description": "Medical devices - Quality management systems",
                "key_requirements": [
                    "Document control procedures",
                    "Management responsibility",
                    "Resource management",
                    "Product realization",
                    "Measurement and improvement",
                    "Risk management integration",
                    "Design controls",
                    "Corrective and preventive actions"
                ]
            },
            "ISO14971": {
                "description": "Medical devices - Application of risk management",
                "key_requirements": [
                    "Risk management process",
                    "Risk analysis and evaluation",
                    "Risk control measures",
                    "Residual risk evaluation",
                    "Risk management report",
                    "Post-production information",
                    "Risk/benefit analysis"
                ]
            },
            "IEC62304": {
                "description": "Medical device software - Software life cycle processes",
                "key_requirements": [
                    "Software development planning",
                    "Software requirements analysis",
                    "Software architectural design",
                    "Software detailed design",
                    "Software implementation",
                    "Software integration and testing",
                    "Software system testing",
                    "Software release"
                ]
            }
        }
        
        self.workflow = self._build_workflow()
        self.token_usage = {"input_tokens": 0, "output_tokens": 0, "total_tokens": 0}
    
    def _build_workflow(self) -> StateGraph:
        """Build the LangGraph workflow following evaluator-optimizer pattern"""
        
        def generator_node(state: ReviewState) -> ReviewState:
            """LLM Generator - Reviews procedure against ISO standards"""
            
            generator_prompt = ChatPromptTemplate.from_template("""
            You are a quality assurance expert specializing in medical device standards.
            
            Review the following procedure against ISO 13485, ISO 14971, and IEC 62304 standards.
            
            Knowledge Base:
            {knowledge_base}
            
            Procedure to Review:
            {procedure}
            
            Provide a comprehensive review including:
            1. Compliance assessment for each relevant standard
            2. Identified gaps or deficiencies
            3. Recommendations for improvement
            4. Risk assessment considerations
            5. Documentation requirements
            
            Be specific and reference relevant clauses where applicable.
            """)
            
            messages = generator_prompt.format_messages(
                knowledge_base=json.dumps(self.knowledge_base, indent=2),
                procedure=state["input_procedure"]
            )
            
            response = self.llm.invoke(messages)
            
            # Track token usage
            if hasattr(response, 'usage_metadata'):
                self.token_usage["input_tokens"] += response.usage_metadata.get("input_tokens", 0)
                self.token_usage["output_tokens"] += response.usage_metadata.get("output_tokens", 0)
                self.token_usage["total_tokens"] += response.usage_metadata.get("total_tokens", 0)
            
            return {
                **state,
                "generated_review": response.content
            }
        
        def evaluator_node(state: ReviewState) -> ReviewState:
            """LLM Evaluator - Evaluates the quality of the review"""
            
            evaluator_prompt = ChatPromptTemplate.from_template("""
            You are a senior quality auditor evaluating a procedure review.
            
            Original Procedure:
            {procedure}
            
            Generated Review:
            {review}
            
            Evaluate this review based on:
            1. Completeness - Does it cover all relevant ISO standards?
            2. Accuracy - Are the references and interpretations correct?
            3. Specificity - Are recommendations specific and actionable?
            4. Risk Assessment - Is risk management adequately addressed?
            5. Documentation - Are documentation requirements clearly stated?
            
            Rate each criterion (1-5 scale) and provide overall assessment.
            If score is below 4 in any area, provide specific feedback for improvement.
            
            Respond in JSON format:
            {{
                "scores": {{
                    "completeness": <score>,
                    "accuracy": <score>,
                    "specificity": <score>,
                    "risk_assessment": <score>,
                    "documentation": <score>
                }},
                "overall_score": <average_score>,
                "passed": <true/false>,
                "feedback": "<detailed feedback for improvement>"
            }}
            """)
            
            messages = evaluator_prompt.format_messages(
                procedure=state["input_procedure"],
                review=state["generated_review"]
            )
            
            response = self.llm.invoke(messages)
            
            # Track token usage
            if hasattr(response, 'usage_metadata'):
                self.token_usage["input_tokens"] += response.usage_metadata.get("input_tokens", 0)
                self.token_usage["output_tokens"] += response.usage_metadata.get("output_tokens", 0)
                self.token_usage["total_tokens"] += response.usage_metadata.get("total_tokens", 0)
            
            try:
                evaluation = json.loads(response.content)
            except:
                # Fallback if JSON parsing fails
                evaluation = {
                    "scores": {"completeness": 3, "accuracy": 3, "specificity": 3, "risk_assessment": 3, "documentation": 3},
                    "overall_score": 3.0,
                    "passed": False,
                    "feedback": "Review needs improvement in multiple areas."
                }
            
            return {
                **state,
                "evaluation_result": evaluation,
                "feedback": evaluation["feedback"]
            }
        
        def decision_node(state: ReviewState) -> str:
            """Decision node to determine if review is acceptable or needs refinement"""
            if state["evaluation_result"]["passed"] or state["iteration_count"] >= state["max_iterations"]:
                return "output"
            else:
                return "generator"
        
        def output_node(state: ReviewState) -> ReviewState:
            """Final output node"""
            return {
                **state,
                "final_output": state["generated_review"]
            }
        
        def increment_iteration(state: ReviewState) -> ReviewState:
            """Increment iteration counter"""
            return {
                **state,
                "iteration_count": state["iteration_count"] + 1
            }
        
        # Build the graph
        workflow = StateGraph(ReviewState)
        
        # Add nodes
        workflow.add_node("generator", generator_node)
        workflow.add_node("evaluator", evaluator_node)
        workflow.add_node("increment", increment_iteration)
        workflow.add_node("output", output_node)
        
        # Add edges
        workflow.add_edge("generator", "evaluator")
        workflow.add_edge("evaluator", "increment")
        workflow.add_conditional_edges(
            "increment",
            decision_node,
            {
                "generator": "generator",
                "output": "output"
            }
        )
        workflow.add_edge("output", END)
        
        # Set entry point
        workflow.set_entry_point("generator")
        
        return workflow.compile()
    
    def review_procedure(self, procedure_text: str, max_iterations: int = 3) -> Dict[str, Any]:
        """Run the quality review workflow"""
        
        initial_state = ReviewState(
            input_procedure=procedure_text,
            generated_review="",
            evaluation_result={},
            feedback="",
            final_output="",
            iteration_count=0,
            max_iterations=max_iterations
        )
        
        # Reset token usage for this review
        self.token_usage = {"input_tokens": 0, "output_tokens": 0, "total_tokens": 0}
        
        # Execute the workflow
        result = self.workflow.invoke(initial_state)
        
        # Add token usage to result
        result["token_usage"] = self.token_usage.copy()
        
        return result
    
    def generate_report(self, review_result: Dict[str, Any], output_filename: str = None) -> str:
        """Generate a DOCX report from the review results"""
        
        if output_filename is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_filename = f"quality_review_report_{timestamp}.docx"
        
        # Create a new document
        doc = Document()
        
        # Add title
        title = doc.add_heading('Quality Procedure Review Report', 0)
        
        # Add metadata
        doc.add_heading('Report Information', level=1)
        info_table = doc.add_table(rows=4, cols=2)
        info_table.style = 'Table Grid'
        
        info_data = [
            ['Generated Date', datetime.now().strftime("%Y-%m-%d %H:%M:%S")],
            ['Standards Applied', 'ISO 13485, ISO 14971, IEC 62304'],
            ['Iterations Completed', str(review_result['iteration_count'])],
            ['Final Score', f"{review_result['evaluation_result'].get('overall_score', 'N/A')}/5"]
        ]
        
        for i, (key, value) in enumerate(info_data):
            info_table.cell(i, 0).text = key
            info_table.cell(i, 1).text = value
        
        # Add evaluation scores
        if 'evaluation_result' in review_result and 'scores' in review_result['evaluation_result']:
            doc.add_heading('Evaluation Scores', level=1)
            scores_table = doc.add_table(rows=6, cols=2)
            scores_table.style = 'Table Grid'
            
            scores = review_result['evaluation_result']['scores']
            score_data = [
                ['Criterion', 'Score (1-5)'],
                ['Completeness', str(scores.get('completeness', 'N/A'))],
                ['Accuracy', str(scores.get('accuracy', 'N/A'))],
                ['Specificity', str(scores.get('specificity', 'N/A'))],
                ['Risk Assessment', str(scores.get('risk_assessment', 'N/A'))],
                ['Documentation', str(scores.get('documentation', 'N/A'))]
            ]
            
            for i, (criterion, score) in enumerate(score_data):
                scores_table.cell(i, 0).text = criterion
                scores_table.cell(i, 1).text = score
                if i == 0:  # Header row
                    for j in range(2):
                        scores_table.cell(i, j).paragraphs[0].runs[0].bold = True
        
        # Add original procedure
        doc.add_heading('Original Procedure', level=1)
        doc.add_paragraph(review_result['input_procedure'])
        
        # Add review findings
        doc.add_heading('Review Findings', level=1)
        doc.add_paragraph(review_result['final_output'])
        
        # Add feedback if available
        if review_result.get('feedback'):
            doc.add_heading('Evaluation Feedback', level=1)
            doc.add_paragraph(review_result['feedback'])
        
        # Add standards reference
        doc.add_heading('Standards Reference', level=1)
        for standard, details in self.knowledge_base.items():
            doc.add_heading(f'{standard}: {details["description"]}', level=2)
            doc.add_paragraph("Key Requirements:")
            for requirement in details["key_requirements"]:
                p = doc.add_paragraph(requirement)
                p.style = 'List Bullet'
        
        # Save the document
        doc.save(output_filename)
        
        return output_filename

# Example usage and testing
def main():
    """Example usage of the Quality Review Agent"""
    
    # Initialize agent (you'll need to provide your OpenAI API key)
    # agent = QualityReviewAgent(openai_api_key="your-api-key-here")
    
    # Example procedure text
    example_procedure = """
    Procedure: Software Update Process
    
    1. Developer creates software update
    2. Technical lead reviews code changes
    3. QA team performs functional testing
    4. Release manager deploys to production
    5. Monitor system for 24 hours post-deployment
    
    Documentation:
    - Code review checklist completed
    - Test results documented
    - Deployment log maintained
    """
    
    print("Quality Procedure Review Agent")
    print("=" * 50)
    print("\nExample procedure:")
    print(example_procedure)
    
    # Uncomment to run with actual API key
    """
    print("\nRunning review workflow...")
    result = agent.review_procedure(example_procedure)
    
    print(f"\nReview completed in {result['iteration_count']} iterations")
    print(f"Final evaluation score: {result['evaluation_result'].get('overall_score', 'N/A')}/5")
    
    # Generate report
    report_filename = agent.generate_report(result)
    print(f"\nReport generated: {report_filename}")
    """

if __name__ == "__main__":
    main()

# Additional utility functions for enhanced functionality

class ProcedureAnalyzer:
    """Additional utilities for procedure analysis"""
    
    @staticmethod
    def extract_procedure_elements(procedure_text: str) -> Dict[str, List[str]]:
        """Extract key elements from procedure text"""
        lines = procedure_text.strip().split('\n')
        
        elements = {
            "steps": [],
            "documentation": [],
            "roles": [],
            "controls": []
        }
        
        current_section = None
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            # Identify sections
            if "documentation" in line.lower():
                current_section = "documentation"
                continue
            elif line.startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.')):
                elements["steps"].append(line)
            elif line.startswith('-') and current_section == "documentation":
                elements["documentation"].append(line[1:].strip())
            
            # Extract roles mentioned
            roles_keywords = ['developer', 'manager', 'lead', 'team', 'reviewer', 'auditor']
            for keyword in roles_keywords:
                if keyword.lower() in line.lower() and keyword not in elements["roles"]:
                    elements["roles"].append(keyword.title())
        
        return elements
    
    @staticmethod
    def check_iso_alignment(procedure_elements: Dict[str, List[str]]) -> Dict[str, bool]:
        """Basic check for ISO alignment"""
        alignment = {
            "has_documentation": len(procedure_elements["documentation"]) > 0,
            "has_defined_roles": len(procedure_elements["roles"]) > 0,
            "has_process_steps": len(procedure_elements["steps"]) > 0,
            "has_controls": "review" in str(procedure_elements).lower() or "approval" in str(procedure_elements).lower()
        }
        return alignment

# Enhanced workflow with procedure analysis
class EnhancedQualityReviewAgent(QualityReviewAgent):
    """Enhanced version with additional analysis capabilities"""
    
    def __init__(self, openai_api_key: str):
        super().__init__(openai_api_key)
        self.analyzer = ProcedureAnalyzer()
    
    def comprehensive_review(self, procedure_text: str, max_iterations: int = 3) -> Dict[str, Any]:
        """Run comprehensive review with additional analysis"""
        
        # Pre-analysis
        procedure_elements = self.analyzer.extract_procedure_elements(procedure_text)
        iso_alignment = self.analyzer.check_iso_alignment(procedure_elements)
        
        # Run main workflow
        review_result = self.review_procedure(procedure_text, max_iterations)
        
        # Add analysis results
        review_result.update({
            "procedure_elements": procedure_elements,
            "iso_alignment": iso_alignment,
            "pre_analysis_score": sum(iso_alignment.values()) / len(iso_alignment)
        })
        
        return review_result
    
    def generate_enhanced_report(self, review_result: Dict[str, Any], output_filename: str = None) -> str:
        """Generate enhanced DOCX report with additional analysis"""
        
        if output_filename is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_filename = f"enhanced_quality_review_{timestamp}.docx"
        
        doc = Document()
        
        # Title and executive summary
        doc.add_heading('Quality Procedure Review Report', 0)
        doc.add_heading('Executive Summary', level=1)
        
        summary = f"""
        This report presents a comprehensive review of the submitted procedure against medical device quality standards.
        The procedure underwent {review_result['iteration_count']} evaluation iterations with a final score of 
        {review_result['evaluation_result'].get('overall_score', 'N/A')}/5.
        """
        doc.add_paragraph(summary.strip())
        
        # Quick assessment table
        doc.add_heading('Quick Assessment', level=1)
        if 'iso_alignment' in review_result:
            alignment_table = doc.add_table(rows=5, cols=2)
            alignment_table.style = 'Table Grid'
            
            alignment_data = [
                ['Assessment Criteria', 'Status'],
                ['Documentation Present', '✓' if review_result['iso_alignment']['has_documentation'] else '✗'],
                ['Roles Defined', '✓' if review_result['iso_alignment']['has_defined_roles'] else '✗'],
                ['Process Steps Clear', '✓' if review_result['iso_alignment']['has_process_steps'] else '✗'],
                ['Controls in Place', '✓' if review_result['iso_alignment']['has_controls'] else '✗']
            ]
            
            for i, (criteria, status) in enumerate(alignment_data):
                alignment_table.cell(i, 0).text = criteria
                alignment_table.cell(i, 1).text = status
                if i == 0:  # Header row
                    for j in range(2):
                        alignment_table.cell(i, j).paragraphs[0].runs[0].bold = True
        
        # Detailed evaluation scores
        doc.add_heading('Detailed Evaluation', level=1)
        if 'evaluation_result' in review_result and 'scores' in review_result['evaluation_result']:
            scores_table = doc.add_table(rows=6, cols=3)
            scores_table.style = 'Table Grid'
            
            scores = review_result['evaluation_result']['scores']
            score_data = [
                ['Criterion', 'Score', 'Status'],
                ['Completeness', str(scores.get('completeness', 'N/A')), 'Pass' if scores.get('completeness', 0) >= 4 else 'Needs Improvement'],
                ['Accuracy', str(scores.get('accuracy', 'N/A')), 'Pass' if scores.get('accuracy', 0) >= 4 else 'Needs Improvement'],
                ['Specificity', str(scores.get('specificity', 'N/A')), 'Pass' if scores.get('specificity', 0) >= 4 else 'Needs Improvement'],
                ['Risk Assessment', str(scores.get('risk_assessment', 'N/A')), 'Pass' if scores.get('risk_assessment', 0) >= 4 else 'Needs Improvement'],
                ['Documentation', str(scores.get('documentation', 'N/A')), 'Pass' if scores.get('documentation', 0) >= 4 else 'Needs Improvement']
            ]
            
            for i, (criterion, score, status) in enumerate(score_data):
                scores_table.cell(i, 0).text = criterion
                scores_table.cell(i, 1).text = score
                scores_table.cell(i, 2).text = status
                if i == 0:  # Header row
                    for j in range(3):
                        scores_table.cell(i, j).paragraphs[0].runs[0].bold = True
        
        # Original procedure
        doc.add_heading('Original Procedure', level=1)
        doc.add_paragraph(review_result['input_procedure'])
        
        # Procedure elements analysis
        if 'procedure_elements' in review_result:
            doc.add_heading('Procedure Analysis', level=1)
            elements = review_result['procedure_elements']
            
            doc.add_heading('Identified Steps', level=2)
            for step in elements['steps']:
                doc.add_paragraph(step, style='List Number')
            
            doc.add_heading('Documentation Requirements', level=2)
            for doc_req in elements['documentation']:
                doc.add_paragraph(doc_req, style='List Bullet')
            
            doc.add_heading('Identified Roles', level=2)
            for role in elements['roles']:
                doc.add_paragraph(role, style='List Bullet')
        
        # Review findings
        doc.add_heading('Detailed Review Findings', level=1)
        doc.add_paragraph(review_result['final_output'])
        
        # Recommendations
        doc.add_heading('Recommendations', level=1)
        if review_result.get('feedback'):
            doc.add_paragraph(review_result['feedback'])
        
        # Standards appendix
        doc.add_page_break()
        doc.add_heading('Standards Reference', level=1)
        for standard, details in self.knowledge_base.items():
            doc.add_heading(f'{standard}: {details["description"]}', level=2)
            doc.add_paragraph("Key Requirements:")
            for requirement in details["key_requirements"]:
                p = doc.add_paragraph(requirement)
                p.style = 'List Bullet'
        
        # Save document
        doc.save(output_filename)
        
        return output_filename

# Example usage script
def run_example():
    """Example script showing how to use the agent"""
    
    print("Quality Procedure Review Agent - Example")
    print("=" * 50)
    
    # Example procedure for software development
    example_procedure = """
    Software Change Control Procedure - SOP-SW-001
    
    Purpose: To ensure all software changes are properly controlled and validated
    
    Procedure Steps:
    1. Developer submits change request with rationale
    2. Technical lead reviews proposed changes for technical merit
    3. Quality assurance reviews for compliance impact
    4. Risk manager assesses safety implications
    5. Configuration manager updates version control
    6. Testing team executes validation protocol
    7. Quality manager approves release
    8. System administrator deploys to production environment
    
    Documentation Requirements:
    - Change request form completed
    - Technical review checklist signed
    - Risk assessment documented
    - Test protocol executed and results recorded
    - Release approval documented
    - Deployment log maintained
    
    Roles and Responsibilities:
    - Developer: Initiates change, implements modifications
    - Technical Lead: Reviews technical aspects
    - QA Team: Validates compliance requirements
    - Risk Manager: Assesses safety implications
    """
    
    print("Sample Procedure:")
    print(example_procedure)
    print("\n" + "=" * 50)
    
    # Instructions for actual usage
    usage_instructions = """
    To use this agent:
    
    1. Install required packages:
       pip install langchain-openai langgraph python-docx
    
    2. Set up your OpenAI API key:
       export OPENAI_API_KEY="your-api-key-here"
       
    3. Initialize and run the agent:
       
       agent = EnhancedQualityReviewAgent(os.getenv("OPENAI_API_KEY"))
       result = agent.comprehensive_review(your_procedure_text)
       report_file = agent.generate_enhanced_report(result)
       
    4. The agent will:
       - Analyze your procedure against ISO standards
       - Use the evaluator-optimizer pattern for quality assurance
       - Generate a comprehensive DOCX report
       - Provide specific recommendations for improvement
    """
    
    print(usage_instructions)

if __name__ == "__main__":
    run_example()